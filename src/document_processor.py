import os
from google.cloud import vision
import pandas as pd
from docx import Document
import fitz
import json
from pathlib import Path
from typing import Dict, List, Any, Optional
import re
from PIL import Image
import pytesseract
import base64
import io
import hashlib
from .ai_client import UnifiedAIClient

class DocumentProcessor:
    def __init__(self, ai_provider: str = "openai", api_key: str = None, model: str = None,
                 google_vision_key_path: str = "key.json", **kwargs):
        """Initialize the document processor with unified AI client and optional Google Vision API"""
        
        # Set defaults based on provider
        if ai_provider.lower() == "gemini":
            default_model = model or "gemini-1.5-flash"
            api_key = api_key or os.getenv('GEMINI_API_KEY')
        else:  # OpenAI
            default_model = model or "gpt-3.5-turbo"
            api_key = api_key or os.getenv('OPENAI_API_KEY')
        
        # Initialize unified AI client
        try:
            self.ai_client = UnifiedAIClient(
                provider=ai_provider,
                api_key=api_key,
                model=default_model,
                **kwargs
            )
            self.ai_provider = ai_provider
            self.model = default_model
        except Exception as e:
            print(f"Error initializing AI client: {e}")
            raise
        
        # Initialize Google Vision API
        if os.path.exists(google_vision_key_path):
            os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = google_vision_key_path
            self.vision_client = vision.ImageAnnotatorClient()
        else:
            print(f"Warning: Google Vision API key file not found at {google_vision_key_path}")
            self.vision_client = None
    
    def extract_patient_info_from_text(self, text: str, use_ai: bool = False) -> Dict[str, str]:
        """Extract patient name and ID from text using AI or regex fallback"""
        # Skip AI extraction during initial ingestion to prevent hanging
        if not use_ai:
            return self._extract_patient_info_regex(text)
            
        try:
            prompt =  f"""
                Extract the patient name and patient ID from the following medical text. 
                Return the result strictly as a **JSON object** with the keys **'patient_name'** and **'patient_id'**.
                Please **do not include any markdown, backticks, or additional text** — just return the **JSON object** as is.

                Here’s the medical text: 
                {text[:2000]}  # Limit text to avoid token limits
                """
            
            # Generate response using unified AI client
            ai_response = self.ai_client.generate_text(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=500
            )
            
            result = ai_response['content']
            
            # Clean the result to avoid parsing issues
            cleaned_result = result.strip().replace("\n", " ").replace("\r", " ")
            
            # Try to parse JSON response
            try:
                patient_info = json.loads(cleaned_result)  # Attempt to decode JSON response
                return {
                    'patient_name': patient_info.get('patient_name', ''),
                    'patient_id': patient_info.get('patient_id', '')
                }
            except json.JSONDecodeError as e: 
                # Fallback to regex extraction if AI fails to provide valid JSON
                return self._extract_patient_info_regex(text)
                    
        except Exception as e:
            return self._extract_patient_info_regex(text)
    
    def _extract_patient_info_regex(self, text: str) -> Dict[str, str]:
        """Fallback method to extract patient info using regex"""
        patient_name = ""
        patient_id = ""
        
        # Clean up the text by removing extra whitespace and newlines
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Common patterns for patient names - more specific and ordered by priority
        name_patterns = [
            r'Patient\s+Name\s*:?\s*([A-Za-z][A-Za-z\s]+?)(?:\s|$|,|\n)',  # Patient Name: John Smith
            r'Name\s*:?\s*([A-Za-z][A-Za-z\s]+?)(?:\s|$|,|\n)',            # Name: John Smith
            r'Patient\s*:?\s*([A-Za-z][A-Za-z\s]+?)(?:\s|$|,|\n)',         # Patient: John Smith
            r'(?:Mr|Mrs|Ms|Dr)\.?\s+([A-Za-z][A-Za-z\s]+?)(?:\s|$|,|\n)',  # Mr. John Smith
        ]
        
        # Common patterns for patient IDs - more specific
        id_patterns = [
            r'Patient\s*ID\s*:?\s*([A-Za-z0-9-]+?)(?:\s|$|\n)',           # PATIENT ID: M0001
            r'Medical\s*Record\s*:?\s*([A-Za-z0-9-]+?)(?:\s|$|\n)',       # Medical Record: 123456
            r'MR\s*:?\s*([A-Za-z0-9-]+?)(?:\s|$|\n)',                     # MR: 123456
            r'Record\s*Number\s*:?\s*([A-Za-z0-9-]+?)(?:\s|$|\n)',        # Record Number: 123456
            r'ID\s*:?\s*([A-Za-z0-9-]+?)(?:\s|$|\n)',                     # ID: M0001 (but not "Image ID")
            r'\b([A-Z]{1,3}[-_]?\d{3,6})\b',                              # Pattern like MT-0006, M0001, etc.
        ]
        
        # Extract patient name
        for pattern in name_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                name_candidate = match.group(1).strip()
                # Filter out common false positives
                if (len(name_candidate) > 1 and 
                    not re.match(r'^\d+$', name_candidate) and  # Not just numbers
                    not name_candidate.lower() in ['id', 'transcription', 'description', 'path', 'image']):
                    patient_name = name_candidate
                    break
        
        # Extract patient ID with better filtering
        for pattern in id_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                id_candidate = match.group(1).strip()
                # Make sure it's not part of "Image ID" or similar
                match_start = match.start()
                preceding_text = text[max(0, match_start-10):match_start].lower()
                
                if (len(id_candidate) > 0 and 
                    'image' not in preceding_text and
                    'img' not in preceding_text and
                    len(id_candidate) < 20):  # Reasonable length limit
                    patient_id = id_candidate
                    break
        
        # Return the extracted patient info
        return {'patient_name': patient_name, 'patient_id': patient_id}
    
    def extract_images_from_page(self, page, doc):
        """Extract images from a PDF page."""
        image_list = []

        # Get all images on this specific page
        image_list_raw = page.get_images(full=True)  # Get images from the page only
        
        # Iterate over the images
        for img_index, img in enumerate(image_list_raw):
            xref = img[0]  # Image reference number (xref)
            
            # Create a Pixmap object from the image reference (xref) using the document (doc)
            base_image = fitz.Pixmap(doc, xref)  # Use the doc object, not page.doc
            
            # If the image is in RGB or RGBA, convert it to RGB for consistency
            if base_image.n - 1 in [3, 4]:  # 3 -> RGB, 4 -> RGBA
                base_image = fitz.Pixmap(fitz.csRGB, base_image)  # Convert to RGB

            # Convert the image to PNG bytes
            image_bytes = base_image.tobytes("png")
            
            # Convert to a PIL Image from the PNG byte data
            image = Image.open(io.BytesIO(image_bytes))
            image_list.append(image)

            # Free memory used by the Pixmap object
            base_image = None

        return image_list
    
    def perform_ocr_on_image(self, image: Image) -> str:
        """Perform OCR on an image and return extracted text."""
        text = pytesseract.image_to_string(image)
        return text

    def process_excel_file(self, file_path: str) -> List[Dict]:
        """Process Excel file and extract patient records"""
        results = []
        
        try:
            # Read Excel file with better error handling
            try:
                df = pd.read_excel(file_path, engine='openpyxl')
            except Exception as e:
                print(f"Error reading Excel file {file_path} with openpyxl: {e}")
                # Try with default engine
                df = pd.read_excel(file_path)
            
            print(f"📊 Processing Excel file: {file_path}")
            print(f"   Shape: {df.shape}, Columns: {list(df.columns)}")
            
            # Check if this file contains patient data
            has_patient_data = any('patient' in str(col).lower() for col in df.columns)
            has_medical_data = any(keyword in str(df.columns).lower() for keyword in ['medical', 'diagnosis', 'treatment', 'symptom', 'condition'])
            
            if not has_patient_data and not has_medical_data:
                print(f"   ⚠️ Skipping file - no patient or medical data found in columns")
                return results
            
            # Limit processing for large files to prevent hanging
            max_rows = 100  # Process maximum 100 rows to prevent hanging
            total_rows = len(df)
            if total_rows > max_rows:
                print(f"   ⚠️ Large file detected ({total_rows} rows), processing first {max_rows} rows only")
                df = df.head(max_rows)
            
            # Process each row with progress tracking
            processed_count = 0
            for index, row in df.iterrows():
                # Skip completely empty rows
                if row.isna().all():
                    continue
                
                # Show progress for large files
                if index % 25 == 0 and index > 0:
                    print(f"   📊 Processed {index}/{len(df)} rows...")
                
                # Skip rows that are just image metadata without patient content
                row_text = ' '.join(str(val) for val in row.dropna().values).lower()
                if not any(keyword in row_text for keyword in ['patient', 'medical', 'diagnosis', 'treatment', 'year', 'old', 'male', 'female']):
                    continue
                    
                # Convert row to text
                text_content = ""
                patient_name = ""
                patient_id = ""
                
                # Extract patient info from column names or data
                for col in df.columns:
                    col_lower = str(col).lower()
                    cell_value = row[col]
                    
                    # Skip NaN values
                    if pd.isna(cell_value):
                        continue
                        
                    cell_str = str(cell_value)
                    
                    # Look for patient name in columns (more specific matching)
                    if ('name' in col_lower and 'patient' in col_lower) or ('patient' in col_lower and 'name' in col_lower):
                        # Make sure it's not just "ID" or similar
                        if (len(cell_str) > 2 and 
                            not re.match(r'^[A-Z]*\d+$', cell_str) and  # Not just alphanumeric codes
                            cell_str.lower() not in ['id', 'transcription', 'description']):
                            patient_name = cell_str
                    # Look for patient ID in columns
                    elif ('id' in col_lower and 'patient' in col_lower and 'image' not in col_lower) or col_lower == 'patient id':
                        # Clean patient ID from the cell
                        clean_id = re.sub(r'\s+', ' ', cell_str.strip())
                        if len(clean_id) > 0 and len(clean_id) < 20:  # Reasonable length
                            patient_id = clean_id
                    
                    # Add all content to text (but skip image-only columns for patient extraction)
                    if 'image' not in col_lower or 'patient' in col_lower:
                        text_content += f"{col}: {cell_str}\n"
                
                # Skip rows with no meaningful content
                if not text_content.strip():
                    continue
                
                # If we couldn't find patient info in columns, try to extract from text
                # Skip AI extraction for image-only records (no patient transcriptions)
                if (not patient_name or not patient_id) and 'patient' in text_content.lower():
                    # Use regex fallback first (faster and more reliable)
                    extracted_info = self._extract_patient_info_regex(text_content)
                    if not patient_name and extracted_info.get('patient_name'):
                        patient_name = extracted_info['patient_name']
                    if not patient_id and extracted_info.get('patient_id'):
                        patient_id = extracted_info['patient_id']
                
                # Generate meaningful patient info for medical data
                if not patient_name:
                    if patient_id:
                        # Generate a patient name based on ID
                        patient_name = f"Patient_{patient_id}"
                    else:
                        # Look for age/gender in transcription to create descriptive name
                        age_match = re.search(r'(\d+)[-\s]*year[-\s]*old', text_content, re.IGNORECASE)
                        gender_match = re.search(r'\b(male|female|man|woman)\b', text_content, re.IGNORECASE)
                        
                        descriptors = []
                        if age_match:
                            descriptors.append(f"{age_match.group(1)}yo")
                        if gender_match:
                            descriptors.append(gender_match.group(1).lower())
                        
                        if descriptors:
                            patient_name = f"Patient_{'_'.join(descriptors)}_Row{index}"
                        else:
                            patient_name = f"Patient_Row_{index}"
                
                if not patient_id:
                    # Generate a patient ID if none found
                    file_prefix = Path(file_path).stem[:3].upper()
                    patient_id = f"{file_prefix}_{index:04d}"
                
                results.append({
                    'patient_name': patient_name,
                    'patient_id': patient_id,
                    'content': text_content.strip(),
                    'metadata': {
                        'source': file_path, 
                        'row_index': index,
                        'total_rows': len(df),
                        'columns': list(df.columns),
                        'extraction_method': 'column_based' if patient_name.startswith('Patient_') else 'extracted'
                    }
                })
        
        except Exception as e:
            print(f"Error processing Excel file {file_path}: {e}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
        
        print(f"✅ Extracted {len(results)} records from {file_path}")
        return results
    
    def process_pdf_file(self, file_path: str) -> Dict:
        """Process PDF file, extract text from pages and images."""
        try:
            doc = fitz.open(file_path)  # Open the PDF file
            content = ""

            # Extract text from PDF pages
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)  # Load each page
                content += page.get_text("text") + "\n"  # Extract text from the page
                
                # Extract images and apply OCR
                images = self.extract_images_from_page(page, doc)
                for img in images:
                    ocr_text = self.perform_ocr_on_image(img)
                    content += ocr_text + "\n"

            # Extract patient info from the extracted content
            patient_info = self.extract_patient_info_from_text(content, use_ai=True)
            
            return {
                'patient_name': patient_info['patient_name'],
                'patient_id': patient_info['patient_id'],
                'content': content.strip(),
                'metadata': {'source': file_path, 'file_type': 'pdf'}
            }

        except Exception as e:
            print(f"Error processing PDF file {file_path}: {e}")
            return {'patient_name': '', 'patient_id': '', 'content': '', 'metadata': {}}
    
    def process_docx_file(self, file_path: str) -> Dict:
        """Process DOCX file and extract content"""
        try:
            doc = Document(file_path)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            # Extract patient info
            patient_info = self.extract_patient_info_from_text(content, use_ai=True)
            
            return {
                'patient_name': patient_info['patient_name'],
                'patient_id': patient_info['patient_id'],
                'content': content.strip(),
                'metadata': {'source': file_path, 'file_type': 'docx'}
            }
        
        except Exception as e:
            print(f"Error processing DOCX file {file_path}: {e}")
            return {'patient_name': '', 'patient_id': '', 'content': '', 'metadata': {}}
    
    def process_text_file(self, file_path: str) -> Dict:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract patient info
            patient_info = self.extract_patient_info_from_text(content, use_ai=True)
            
            return {
                'patient_name': patient_info['patient_name'],
                'patient_id': patient_info['patient_id'],
                'content': content.strip(),
                'metadata': {'source': file_path, 'file_type': 'txt'}
            }
        
        except Exception as e:
            print(f"Error processing text file {file_path}: {e}")
            return {'patient_name': '', 'patient_id': '', 'content': '', 'metadata': {}}
    
    def analyze_image_with_ai_vision(self, image_path: str, patient_name: str = "", 
                                   patient_id: str = "", clinical_context: str = "") -> str:
        """Analyze medical image using unified AI vision capabilities"""
        try:
            # Determine image category
            image_category = self._get_image_category(image_path)
            
            # Create comprehensive prompt for medical image analysis
            prompt = f"""You are a medical AI assistant analyzing a medical image. Please provide a detailed analysis.

PATIENT INFORMATION:
- Name: {patient_name}
- Patient ID: {patient_id}
- Image Type: {image_category}

CLINICAL CONTEXT:
{clinical_context if clinical_context else "No additional clinical context provided"}

Please analyze this medical image and provide:

1. **Image Type & Quality Assessment:**
   - What type of medical imaging is this?
   - Is the image quality adequate for diagnosis?

2. **Anatomical Observations:**
   - What anatomical structures are visible?
   - Any obvious abnormalities or variations?

3. **Medical Findings:**
   - Describe any pathological findings
   - Note any areas of concern
   - Identify normal vs abnormal features

4. **Clinical Correlation:**
   - How do these findings relate to the patient's clinical context?
   - What additional imaging or tests might be recommended?

5. **Summary:**
   - Provide a concise summary of key findings
   - Note any limitations in the analysis

Please be thorough but acknowledge that this is an AI analysis and should be reviewed by qualified medical professionals."""

            # Use unified AI client for image analysis
            if self.ai_client.supports_vision():
                ai_response = self.ai_client.analyze_image(
                    image_path=image_path,
                    prompt=prompt,
                    max_tokens=1500,
                    temperature=0.1
                )
                
                analysis = ai_response['content']
                
                # Format the response with patient info
                formatted_analysis = f"""
PATIENT: {patient_name} (ID: {patient_id})
IMAGE TYPE: {image_category}
AI VISUAL ANALYSIS ({self.ai_provider.upper()}):

{analysis}

---
Note: This analysis was generated using {self.ai_provider.upper()} vision technology and should be reviewed by qualified medical professionals.
"""
                return formatted_analysis
                
            else:
                return f"Vision analysis not supported with current AI configuration for patient {patient_name}"
                
        except Exception as e:
            print(f"Error analyzing image with AI Vision: {e}")
            return f"Error analyzing image with AI Vision: {str(e)}"
    
    def transcribe_image_google_vision(self, image_path: str) -> str:
        """Transcribe image using Google Vision API"""
        if not self.vision_client:
            return "Google Vision API not configured"
        
        try:
            with io.open(image_path, 'rb') as image_file:
                content = image_file.read()
            
            image = vision.Image(content=content)
            response = self.vision_client.text_detection(image=image)
            texts = response.text_annotations
            
            if texts:
                return texts[0].description
            else:
                return "No text detected in image"
        
        except Exception as e:
            print(f"Error transcribing image {image_path}: {e}")
            return f"Error: {str(e)}"
    
    def analyze_image_content(self, image_path: str, transcription: str) -> Dict:
        """Analyze image content and extract medical information"""
        try:
            # Get image category from path
            image_category = self._get_image_category(image_path)
            
            # Use AI to analyze the transcription for medical content
            if transcription and transcription != "No text detected in image":
                prompt = f"""
                Analyze this medical image transcription and extract relevant information:
                
                Image Category: {image_category}
                Transcription: {transcription}
                
                Extract:
                1. Patient name and ID if mentioned
                2. Medical findings
                3. Diagnosis or observations
                4. Any other relevant medical information
                
                Return as JSON with keys: patient_name, patient_id, findings, diagnosis, notes
                """
                
                # Generate analysis using unified AI client
                ai_response = self.ai_client.generate_text(
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                    max_tokens=1000
                )
                
                result = ai_response['content']
                
                try:
                    analysis = json.loads(result)
                    return {
                        'patient_name': analysis.get('patient_name', ''),
                        'patient_id': analysis.get('patient_id', ''),
                        'transcription': transcription,
                        'category': image_category,
                        'analysis': analysis,
                        'metadata': {'source': image_path, 'processed_with_ai': True}
                    }
                except json.JSONDecodeError:
                    # Fallback if JSON parsing fails
                    patient_info = self.extract_patient_info_from_text(transcription)
                    return {
                        'patient_name': patient_info['patient_name'],
                        'patient_id': patient_info['patient_id'],
                        'transcription': transcription,
                        'category': image_category,
                        'analysis': {'notes': result},
                        'metadata': {'source': image_path, 'processed_with_ai': True}
                    }
            else:
                return {
                    'patient_name': '',
                    'patient_id': '',
                    'transcription': transcription,
                    'category': image_category,
                    'analysis': {},
                    'metadata': {'source': image_path, 'processed_with_ai': False}
                }
        
        except Exception as e:
            print(f"Error analyzing image content: {e}")
            return {
                'patient_name': '',
                'patient_id': '',
                'transcription': transcription,
                'category': self._get_image_category(image_path),
                'analysis': {'error': str(e)},
                'metadata': {'source': image_path, 'error': True}
            }
    
    def _get_image_category(self, image_path: str) -> str:
        """Extract image category from file path"""
        path_parts = Path(image_path).parts
        
        # Look for medical imaging categories in path (case-insensitive)
        medical_keywords = {
            'ct': 'CT Scan',
            'mri': 'MRI',
            'xray': 'X-Ray',
            'cxr': 'Chest X-Ray',
            'ultrasound': 'Ultrasound',
            'echo': 'Echocardiogram',
            'mammogram': 'Mammogram',
            'abdomen': 'Abdominal Imaging',
            'chest': 'Chest Imaging',
            'head': 'Head Imaging',
            'brain': 'Brain Imaging',
            'spine': 'Spine Imaging',
            'hand': 'Hand Imaging',
            'foot': 'Foot Imaging',
            'breast': 'Breast Imaging',
            'cardiac': 'Cardiac Imaging',
            'lung': 'Lung Imaging',
            'liver': 'Liver Imaging',
            'kidney': 'Kidney Imaging'
        }
        
        # Check each part of the path for medical keywords
        for part in path_parts:
            part_lower = part.lower()
            for keyword, category in medical_keywords.items():
                if keyword in part_lower:
                    return category
        
        # Try to extract from common folder naming patterns
        for part in path_parts:
            part_lower = part.lower()
            # Look for patterns like "CT_Scans", "MRI_Images", etc.
            if 'ct' in part_lower and ('scan' in part_lower or 'image' in part_lower):
                return 'CT Scan'
            elif 'mri' in part_lower:
                return 'MRI'
            elif 'xray' in part_lower or 'x-ray' in part_lower:
                return 'X-Ray'
            elif 'ultrasound' in part_lower or 'us' in part_lower:
                return 'Ultrasound'
        
        # If no specific category found, use the parent folder name
        if len(path_parts) > 1:
            parent_folder = path_parts[-2]  # Get immediate parent folder
            return f"Medical Image - {parent_folder}"
        
        return 'Medical Image - Uncategorized'
    
    def enhance_text_with_ai(self, text: str, file_type: str = "document") -> str:
        """Enhance and structure text using AI for better searchability"""
        try:
            prompt = f"""
            Please analyze and enhance this medical {file_type} text for better organization and searchability.
            Extract key medical information, standardize terminology, and create a structured summary.
            
            Original text: {text[:3000]}  # Limit to avoid token limits
            
            Please format the response with clear sections like:
            - Patient Information
            - Medical History
            - Symptoms/Complaints
            - Diagnosis
            - Treatment Plan
            - Follow-up
            - Other Notes
            
            Only include sections that have relevant information.
            """
            
            # Generate enhanced text using unified AI client
            ai_response = self.ai_client.generate_text(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=2000
            )
            
            enhanced_text = ai_response['content']
            return enhanced_text
        
        except Exception as e:
            print(f"Error enhancing text with AI: {e}")
            return text  # Return original text if enhancement fails
    
    def process_file(self, file_path: str) -> List[Dict]:
        """Process any supported file type and return extracted data"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in ['.xlsx', '.xls']:
            return self.process_excel_file(file_path)
        elif file_ext == '.pdf':
            result = self.process_pdf_file(file_path)
            return [result] if result['content'] else []
        elif file_ext == '.docx':
            result = self.process_docx_file(file_path)
            return [result] if result['content'] else []
        elif file_ext == '.txt':
            result = self.process_text_file(file_path)
            return [result] if result['content'] else []
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            transcription = self.transcribe_image_google_vision(file_path)
            result = self.analyze_image_content(file_path, transcription)
            return [result] if result else []
        else:
            print(f"Unsupported file type: {file_ext}")
            return []
    
    def enhance_image_analysis_with_patient_context(self, transcription: str, image_category: str, 
                                                   patient_name: str, patient_id: str, clinical_context: str = "") -> str:
        """Enhance image transcription with known patient context and clinical information"""
        try:
            prompt = f"""
            You are analyzing a medical image for a known patient. Provide a comprehensive analysis that combines:
            1. The image transcription
            2. The known patient information
            3. Clinical context from their medical records
            
            PATIENT INFORMATION:
            Name: {patient_name}
            Patient ID: {patient_id}
            
            IMAGE DETAILS:
            Category: {image_category}
            Transcription: {transcription}
            
            CLINICAL CONTEXT:
            {clinical_context if clinical_context else "No additional clinical context available"}
            
            Please provide a structured analysis including:
            - Patient identification
            - Image type and findings
            - Relevant observations from the transcription
            - Correlation with clinical context (if available)
            - Key medical findings or abnormalities noted
            
            Format as a clear, structured medical report.
            """
            
            # Generate enhanced analysis using unified AI client
            ai_response = self.ai_client.generate_text(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=1500
            )
            
            enhanced_analysis = ai_response['content']
            
            # Prepend patient info for easy identification
            final_transcription = f"PATIENT: {patient_name} (ID: {patient_id})\n"
            final_transcription += f"IMAGE TYPE: {image_category}\n"
            final_transcription += "=" * 50 + "\n\n"
            final_transcription += enhanced_analysis
            final_transcription += "\n\n" + "=" * 50
            final_transcription += f"\nORIGINAL TRANSCRIPTION: {transcription}"
            
            return final_transcription
            
        except Exception as e:
            print(f"Error enhancing image analysis with patient context: {e}")
            # Fallback: at least include patient info
            fallback_transcription = f"PATIENT: {patient_name} (ID: {patient_id})\n"
            fallback_transcription += f"IMAGE TYPE: {image_category}\n"
            fallback_transcription += "=" * 50 + "\n\n"
            fallback_transcription += f"Image transcription: {transcription}\n\n"
            if clinical_context:
                fallback_transcription += f"Clinical Context: {clinical_context}\n\n"
            fallback_transcription += "Note: AI enhancement failed, showing basic transcription with patient context."
            
            return fallback_transcription

    def transcribe_and_analyze_image(self, image_path: str, patient_name: str = "", 
                                   patient_id: str = "", clinical_context: str = "") -> Dict[str, str]:
        """Comprehensive image analysis using both text extraction and visual analysis"""
        result = {
            'patient_name': patient_name,
            'patient_id': patient_id,
            'image_category': self._get_image_category(image_path),
            'text_extraction': '',
            'visual_analysis': '',
            'combined_analysis': '',
            'processing_method': ''
        }
        
        # Try Google Vision for text extraction first
        if self.vision_client:
            try:
                result['text_extraction'] = self.transcribe_image_google_vision(image_path)
                result['processing_method'] = f'Google Vision + {self.ai_provider.upper()} Analysis'
            except Exception as e:
                print(f"Google Vision failed: {e}")
                result['text_extraction'] = "Text extraction failed"
        else:
            result['text_extraction'] = "Google Vision API not configured"
        
        # Try AI Vision for comprehensive analysis
        try:
            result['visual_analysis'] = self.analyze_image_with_ai_vision(
                image_path, patient_name, patient_id, clinical_context
            )
            if result['processing_method'] == '':
                result['processing_method'] = f'{self.ai_provider.upper()} Vision Analysis Only'
        except Exception as e:
            print(f"AI Vision failed: {e}")
            result['visual_analysis'] = f"Visual analysis failed: {str(e)}"
            if result['processing_method'] == '':
                result['processing_method'] = 'No analysis available'
        
        # Create combined analysis
        if result['visual_analysis'] and "Error" not in result['visual_analysis']:
            # Visual analysis is primary
            result['combined_analysis'] = result['visual_analysis']
            if result['text_extraction'] and result['text_extraction'] not in ["No text detected in image", "Google Vision API not configured", "Text extraction failed"]:
                result['combined_analysis'] += f"\n\nEXTRACTED TEXT:\n{result['text_extraction']}"
        elif result['text_extraction'] and result['text_extraction'] not in ["No text detected in image", "Google Vision API not configured", "Text extraction failed"]:
            # Fallback to text-based analysis
            result['combined_analysis'] = f"Patient: {patient_name} (ID: {patient_id})\nImage Category: {result['image_category']}\n\nExtracted Text Analysis:\n{result['text_extraction']}"
            if result['processing_method'] == '':
                result['processing_method'] = 'Text extraction only'
        else:
            # Minimal fallback
            result['combined_analysis'] = f"Patient: {patient_name} (ID: {patient_id})\nImage Category: {result['image_category']}\nNote: Image analysis not available - please check API configurations"
            result['processing_method'] = 'Basic metadata only'
        
        return result 